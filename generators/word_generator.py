"""
Word Document Generator for Equity Research Reports.

This module generates professional Word documents (.docx) following
the assignment template structure with proper formatting, tables,
and styling.
"""

import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, Any
import pandas as pd

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from agents.state import EquityResearchState
from utils.logger import logger


def generate_word_report(state: EquityResearchState, output_dir: str = "output") -> str:
    """
    Generate comprehensive Word document report.
    
    Creates a professionally formatted .docx file following the
    assignment template with all analysis sections.
    
    Args:
        state: Complete EquityResearchState with all data and analysis
        output_dir: Directory to save the report (default: "output")
    
    Returns:
        str: Path to generated Word document
    
    Example:
        >>> from agents import run_research_workflow
        >>> state = run_research_workflow("RELIANCE")
        >>> path = generate_word_report(state)
        >>> print(f"Report saved: {path}")
    """
    logger.info(f"\n{'='*70}")
    logger.info(f"📄 GENERATING WORD REPORT: {state['company_name']} ({state['ticker']})")
    logger.info(f"{'='*70}\n")
    
    # Create output directory
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Initialize document
    doc = Document()
    
    # Configure document properties
    _setup_document_styles(doc)
    
    # Get report date
    report_date = datetime.now().strftime("%B %d, %Y")
    
    # 1. Cover Page
    logger.info("📝 Step 1/10: Creating cover page...")
    _add_cover_page(doc, state, report_date)
    
    # 2. Table of Contents (placeholder)
    logger.info("📝 Step 2/10: Adding table of contents...")
    _add_table_of_contents(doc)
    
    # 3. Executive Summary
    logger.info("📝 Step 3/10: Adding executive summary...")
    _add_executive_summary(doc, state)
    
    # 4. Company Overview
    logger.info("📝 Step 4/10: Adding company overview...")
    _add_company_overview(doc, state)
    
    # 5. Financial Analysis
    logger.info("📝 Step 5/10: Adding financial analysis...")
    _add_financial_analysis(doc, state)
    
    # 6. Valuation Analysis
    logger.info("📝 Step 6/10: Adding valuation analysis...")
    _add_valuation_analysis(doc, state)
    
    # 7. Risk Analysis
    logger.info("📝 Step 7/10: Adding risk analysis...")
    _add_risk_analysis(doc, state)
    
    # 8. Recent Developments
    logger.info("📝 Step 8/10: Adding recent developments...")
    _add_recent_developments(doc, state)
    
    # 9. Investment Recommendation
    logger.info("📝 Step 9/10: Adding investment recommendation...")
    _add_investment_recommendation(doc, state)
    
    # 10. Appendix (Financial Data Tables)
    logger.info("📝 Step 10/10: Adding appendix...")
    _add_appendix(doc, state)
    
    # Save document
    filename = f"Equity_Research_{state['ticker']}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = output_path / filename
    doc.save(str(filepath))
    
    logger.success(f"✅ Word report generated: {filepath}")
    logger.info(f"   File size: {filepath.stat().st_size / 1024:.2f} KB")
    
    return str(filepath)


def _setup_document_styles(doc: Document):
    """Setup custom styles for the document."""
    # Title style
    if 'Report Title' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 - Section Headers
    heading1 = doc.styles['Heading 1']
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(6)
    
    # Heading 2 - Subsection Headers
    heading2 = doc.styles['Heading 2']
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 76, 153)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)


def _add_cover_page(doc: Document, state: EquityResearchState, report_date: str):
    """Add professional cover page."""
    # Title
    title = doc.add_paragraph(style='Report Title')
    title.add_run("EQUITY RESEARCH REPORT")
    title.paragraph_format.space_after = Pt(24)
    
    # Company name
    company = doc.add_paragraph()
    company_run = company.add_run(state.get('company_name', state['ticker']))
    company_run.font.size = Pt(20)
    company_run.font.bold = True
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.paragraph_format.space_after = Pt(12)
    
    # Ticker
    ticker = doc.add_paragraph()
    ticker_run = ticker.add_run(f"({state['ticker']})")
    ticker_run.font.size = Pt(16)
    ticker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ticker.paragraph_format.space_after = Pt(36)
    
    # Key metrics table
    company_info = state.get('company_info', {})
    stock_prices = state.get('stock_prices')
    
    metrics_para = doc.add_paragraph()
    metrics_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    current_price = stock_prices['Close'].iloc[-1] if stock_prices is not None and not stock_prices.empty else 0
    market_cap = company_info.get('marketCap', 0) / 1e9 if company_info.get('marketCap') else 0
    
    metrics_text = f"""
Current Price: ₹{current_price:.2f}
Market Cap: ₹{market_cap:.2f}B
Sector: {company_info.get('sector', 'N/A')}
Industry: {company_info.get('industry', 'N/A')}
    """
    
    metrics_run = metrics_para.add_run(metrics_text.strip())
    metrics_run.font.size = Pt(12)
    metrics_para.paragraph_format.space_after = Pt(48)
    
    # Recommendation
    if state.get('valuation_recommendation'):
        rec_para = doc.add_paragraph()
        rec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rec_run = rec_para.add_run(f"RECOMMENDATION: {state['valuation_recommendation']}")
        rec_run.font.size = Pt(14)
        rec_run.font.bold = True
        
        # Color code recommendation
        if 'Buy' in state['valuation_recommendation']:
            rec_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif 'Sell' in state['valuation_recommendation']:
            rec_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        else:
            rec_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        
        rec_para.paragraph_format.space_after = Pt(48)
    
    # Report date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Report Date: {report_date}")
    date_run.font.size = Pt(12)
    date_para.paragraph_format.space_after = Pt(24)
    
    # Disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer.add_run("Generated by Automated Equity Research System")
    disclaimer_run.font.size = Pt(10)
    disclaimer_run.font.italic = True
    disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Page break
    doc.add_page_break()


def _add_table_of_contents(doc: Document):
    """Add table of contents."""
    doc.add_heading('Table of Contents', 0)
    
    toc_items = [
        "1. Executive Summary",
        "2. Company Overview",
        "3. Financial Analysis",
        "   3.1 Liquidity Analysis",
        "   3.2 Efficiency Analysis",
        "   3.3 Solvency Analysis",
        "   3.4 Profitability Analysis",
        "4. Valuation Analysis",
        "   4.1 Beta and Risk Profile",
        "   4.2 Cost of Equity (CAPM)",
        "   4.3 Dividend Discount Model (DDM)",
        "5. Risk Analysis",
        "6. Recent Developments",
        "7. Investment Recommendation",
        "8. Appendix - Financial Data Tables"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')
    
    doc.add_page_break()


def _add_executive_summary(doc: Document, state: EquityResearchState):
    """Add executive summary section."""
    doc.add_heading('1. Executive Summary', 1)
    
    summary_text = state.get('executive_summary', 
        "[Executive summary will be generated when LLM is configured. "
        "This section provides a concise overview of the company, key findings, "
        "and investment recommendation.]")
    
    doc.add_paragraph(summary_text)
    doc.add_paragraph()  # Spacing


def _add_company_overview(doc: Document, state: EquityResearchState):
    """Add company overview section."""
    doc.add_heading('2. Company Overview', 1)
    
    company_info = state.get('company_info', {})
    
    # Company description
    if state.get('company_overview_text'):
        doc.add_paragraph(state['company_overview_text'])
    else:
        # Fallback to basic info
        if company_info.get('longBusinessSummary'):
            doc.add_paragraph(company_info['longBusinessSummary'][:500] + "...")
        else:
            doc.add_paragraph("[Company overview will be generated when LLM is configured.]")
    
    # Company details table
    doc.add_heading('Company Details', 2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    details = [
        ('Sector', company_info.get('sector', 'N/A')),
        ('Industry', company_info.get('industry', 'N/A')),
        ('Employees', f"{company_info.get('fullTimeEmployees', 'N/A'):,}" if company_info.get('fullTimeEmployees') else 'N/A'),
        ('Website', company_info.get('website', 'N/A')),
        ('Country', company_info.get('country', 'India')),
        ('Exchange', 'NSE (National Stock Exchange of India)')
    ]
    
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()


def _add_financial_analysis(doc: Document, state: EquityResearchState):
    """Add financial analysis section."""
    doc.add_heading('3. Financial Analysis', 1)
    
    # LLM-generated analysis
    if state.get('financial_analysis_text'):
        doc.add_paragraph(state['financial_analysis_text'])
    else:
        doc.add_paragraph("[Financial analysis commentary will be generated when LLM is configured.]")
    
    # Financial ratios tables
    ratios = state.get('ratios', {})
    
    if ratios:
        # Liquidity Ratios
        doc.add_heading('3.1 Liquidity Ratios', 2)
        _add_ratio_table(doc, ratios.get('liquidity', {}))
        
        # Efficiency Ratios
        doc.add_heading('3.2 Efficiency Ratios', 2)
        _add_ratio_table(doc, ratios.get('efficiency', {}))
        
        # Solvency Ratios
        doc.add_heading('3.3 Solvency/Leverage Ratios', 2)
        _add_ratio_table(doc, ratios.get('solvency', {}))
        
        # Profitability Ratios
        doc.add_heading('3.4 Profitability Ratios', 2)
        _add_ratio_table(doc, ratios.get('profitability', {}))
    
    doc.add_paragraph()


def _add_ratio_table(doc: Document, ratios: Dict[str, float]):
    """Add a table of financial ratios."""
    if not ratios:
        doc.add_paragraph("Data not available", style='Intense Quote')
        return
    
    table = doc.add_table(rows=len(ratios) + 1, cols=2)
    table.style = 'Light List Accent 1'
    
    # Header
    table.rows[0].cells[0].text = "Ratio"
    table.rows[0].cells[1].text = "Value"
    
    # Make header bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    for i, (name, value) in enumerate(ratios.items(), 1):
        formatted_name = name.replace('_', ' ').title()
        table.rows[i].cells[0].text = formatted_name
        
        if value is not None:
            # Format based on ratio type
            if 'margin' in name or 'return' in name:
                table.rows[i].cells[1].text = f"{value:.2f}%"
            else:
                table.rows[i].cells[1].text = f"{value:.2f}"
        else:
            table.rows[i].cells[1].text = "N/A"
    
    doc.add_paragraph()


def _add_valuation_analysis(doc: Document, state: EquityResearchState):
    """Add valuation analysis section."""
    doc.add_heading('4. Valuation Analysis', 1)
    
    # LLM-generated valuation commentary
    if state.get('valuation_text'):
        doc.add_paragraph(state['valuation_text'])
    else:
        doc.add_paragraph("[Valuation analysis will be generated when LLM is configured.]")
    
    # Beta & Risk Metrics
    doc.add_heading('4.1 Beta and Risk Profile', 2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    beta = state.get('beta', 0)
    correlation = state.get('correlation_with_market', 0)
    
    risk_data = [
        ('Beta (vs NIFTY 50)', f"{beta:.3f}" if beta else 'N/A'),
        ('Interpretation', 'Aggressive' if beta and beta > 1 else 'Defensive' if beta else 'N/A'),
        ('Correlation with Market', f"{correlation:.3f}" if correlation else 'N/A'),
        ('Systematic Risk', 'Above Market' if beta and beta > 1 else 'Below Market' if beta else 'N/A')
    ]
    
    for i, (label, value) in enumerate(risk_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    # CAPM
    doc.add_heading('4.2 Cost of Equity (CAPM)', 2)
    
    cost_of_equity = state.get('cost_of_equity', 0)
    
    from config.settings import RISK_FREE_RATE, EXPECTED_MARKET_RETURN
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    capm_data = [
        ('Risk-Free Rate (Indian G-Sec)', f"{RISK_FREE_RATE:.2%}"),
        ('Expected Market Return (NIFTY 50)', f"{EXPECTED_MARKET_RETURN:.2%}"),
        ('Beta', f"{beta:.3f}" if beta else 'N/A'),
        ('Cost of Equity', f"{cost_of_equity:.2%}" if cost_of_equity else 'N/A')
    ]
    
    for i, (label, value) in enumerate(capm_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    # DDM Valuation
    doc.add_heading('4.3 Dividend Discount Model (DDM)', 2)
    
    ddm = state.get('ddm_valuation', {})
    
    if ddm and ddm.get('applicable'):
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        stock_prices = state.get('stock_prices')
        current_price = stock_prices['Close'].iloc[-1] if stock_prices is not None and not stock_prices.empty else 0
        
        ddm_data = [
            ('Current Dividend (D0)', f"₹{ddm.get('d0_current_dividend', 0):.2f}"),
            ('Next Dividend (D1)', f"₹{ddm.get('d1_next_dividend', 0):.2f}"),
            ('Growth Rate', f"{ddm.get('growth_rate', 0):.2%}"),
            ('Fair Value', f"₹{ddm.get('fair_value', 0):.2f}"),
            ('Current Price', f"₹{current_price:.2f}"),
            ('Upside/Downside', f"{ddm.get('upside_downside', 0):.1%}")
        ]
        
        for i, (label, value) in enumerate(ddm_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
    else:
        reason = ddm.get('reason', 'DDM not applicable')
        doc.add_paragraph(f"DDM Valuation: {reason}", style='Intense Quote')
    
    doc.add_paragraph()


def _add_risk_analysis(doc: Document, state: EquityResearchState):
    """Add risk analysis section."""
    doc.add_heading('5. Risk Analysis', 1)
    
    if state.get('risk_analysis_text'):
        doc.add_paragraph(state['risk_analysis_text'])
    else:
        doc.add_paragraph("[Risk analysis will be generated when LLM is configured.]")
        
        # Add basic risk metrics table as fallback
        doc.add_heading('Key Risk Metrics', 2)
        
        ratios = state.get('ratios', {})
        beta = state.get('beta', 0)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        risk_metrics = [
            ('Beta (Market Risk)', f"{beta:.3f}" if beta else 'N/A'),
            ('Debt to Equity', f"{ratios.get('solvency', {}).get('debt_to_equity', 'N/A'):.2f}" if ratios.get('solvency', {}).get('debt_to_equity') else 'N/A'),
            ('Interest Coverage', f"{ratios.get('solvency', {}).get('interest_coverage', 'N/A'):.2f}" if ratios.get('solvency', {}).get('interest_coverage') else 'N/A'),
            ('Current Ratio', f"{ratios.get('liquidity', {}).get('current_ratio', 'N/A'):.2f}" if ratios.get('liquidity', {}).get('current_ratio') else 'N/A')
        ]
        
        for i, (label, value) in enumerate(risk_metrics):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()


def _add_recent_developments(doc: Document, state: EquityResearchState):
    """Add recent developments section."""
    doc.add_heading('6. Recent Developments', 1)
    
    if state.get('recent_developments_text'):
        doc.add_paragraph(state['recent_developments_text'])
    else:
        # Fallback to news summary
        news = state.get('news')
        if news is not None and not news.empty:
            doc.add_paragraph("Recent news highlights:")
            
            for _, row in news.head(10).iterrows():
                date = row['published'].strftime('%Y-%m-%d')
                title = row['title']
                para = doc.add_paragraph(f"• [{date}] {title}", style='List Bullet')
        else:
            doc.add_paragraph("[Recent developments will be included when news data is available.]")
    
    doc.add_paragraph()


def _add_investment_recommendation(doc: Document, state: EquityResearchState):
    """Add investment recommendation section."""
    doc.add_heading('7. Investment Recommendation', 1)
    
    if state.get('final_recommendation_text'):
        doc.add_paragraph(state['final_recommendation_text'])
    else:
        # Fallback to basic recommendation
        recommendation = state.get('valuation_recommendation', 'N/A')
        
        doc.add_paragraph(f"Based on our comprehensive analysis, our recommendation is:")
        
        rec_para = doc.add_paragraph()
        rec_run = rec_para.add_run(f"\n{recommendation}\n")
        rec_run.font.size = Pt(14)
        rec_run.font.bold = True
        
        if 'Buy' in recommendation:
            rec_run.font.color.rgb = RGBColor(0, 128, 0)
        elif 'Sell' in recommendation:
            rec_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            rec_run.font.color.rgb = RGBColor(255, 165, 0)
        
        # Add basic rationale
        ddm = state.get('ddm_valuation', {})
        if ddm and ddm.get('applicable'):
            fair_value = ddm.get('fair_value', 0)
            stock_prices = state.get('stock_prices')
            current_price = stock_prices['Close'].iloc[-1] if stock_prices is not None and not stock_prices.empty else 0
            upside = ddm.get('upside_downside', 0)
            
            doc.add_paragraph(
                f"Our DDM valuation indicates a fair value of ₹{fair_value:.2f} "
                f"compared to the current price of ₹{current_price:.2f}, "
                f"representing a {upside:.1%} {'upside' if upside > 0 else 'downside'}."
            )
    
    doc.add_paragraph()


def _add_appendix(doc: Document, state: EquityResearchState):
    """Add appendix with detailed financial data."""
    doc.add_page_break()
    doc.add_heading('8. Appendix - Financial Data Tables', 1)
    
    # Financial Statements Summary
    financial_statements = state.get('financial_statements', {})
    
    if financial_statements:
        # Income Statement
        doc.add_heading('A. Income Statement Summary', 2)
        income = financial_statements.get('income_statement')
        if income is not None and not income.empty:
            _add_dataframe_table(doc, income.head(5), "Income Statement (₹ Crores)")
        else:
            doc.add_paragraph("Income statement data not available")
        
        # Balance Sheet
        doc.add_heading('B. Balance Sheet Summary', 2)
        balance = financial_statements.get('balance_sheet')
        if balance is not None and not balance.empty:
            _add_dataframe_table(doc, balance.head(5), "Balance Sheet (₹ Crores)")
        else:
            doc.add_paragraph("Balance sheet data not available")
        
        # Cash Flow
        doc.add_heading('C. Cash Flow Summary', 2)
        cashflow = financial_statements.get('cash_flow')
        if cashflow is not None and not cashflow.empty:
            _add_dataframe_table(doc, cashflow.head(5), "Cash Flow Statement (₹ Crores)")
        else:
            doc.add_paragraph("Cash flow data not available")


def _add_dataframe_table(doc: Document, df: pd.DataFrame, caption: str):
    """Add a pandas DataFrame as a Word table."""
    if df.empty:
        doc.add_paragraph("No data available")
        return
    
    # Add caption
    doc.add_paragraph(caption, style='Intense Quote')
    
    # Create table
    table = doc.add_table(rows=min(len(df) + 1, 6), cols=min(len(df.columns) + 1, 6))
    table.style = 'Light Grid Accent 1'
    
    # Headers (column dates/periods)
    table.rows[0].cells[0].text = "Item"
    for j, col in enumerate(df.columns[:5], 1):
        if hasattr(col, 'strftime'):
            table.rows[0].cells[j].text = col.strftime('%Y-%m-%d')
        else:
            table.rows[0].cells[j].text = str(col)
    
    # Data rows
    for i, (idx, row) in enumerate(df.head(5).iterrows(), 1):
        table.rows[i].cells[0].text = str(idx)
        for j, val in enumerate(row[:5], 1):
            if pd.notna(val):
                # Format numbers
                try:
                    if abs(val) > 1e6:
                        table.rows[i].cells[j].text = f"{val/1e7:.2f}Cr"
                    else:
                        table.rows[i].cells[j].text = f"{val:.2f}"
                except:
                    table.rows[i].cells[j].text = str(val)
            else:
                table.rows[i].cells[j].text = "N/A"
    
    doc.add_paragraph()


if __name__ == "__main__":
    """Test Word document generation."""
    print("Testing Word Document Generator...")
    
    from agents.state import create_initial_state
    from agents.nodes import collect_data_node, analyze_node
    
    test_ticker = "RELIANCE"
    
    try:
        print(f"\n🧪 Testing with {test_ticker}...")
        
        # Collect data and analyze
        print("📊 Collecting data...")
        state = create_initial_state(test_ticker, "Reliance Industries")
        data_updates = collect_data_node(state)
        state.update(data_updates)
        
        print("📈 Analyzing...")
        analysis_updates = analyze_node(state)
        state.update(analysis_updates)
        
        # Generate Word report
        print("\n📄 Generating Word document...")
        filepath = generate_word_report(state)
        
        print(f"\n✅ Test completed!")
        print(f"   Report saved: {filepath}")
        
    except Exception as e:
        print(f"\n❌ Test failed: {e}")
        import traceback
        traceback.print_exc()

